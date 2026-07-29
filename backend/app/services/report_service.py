"""报告导出服务（T4.17，详细设计 1.11）。

三模板：topic_deep（议题深度报告）/ compare_brief（跨国对比简报）/ periodic_weekly（周期监测周报）。
- 时间窗上限 90 天预检（超限 CODE_PARAM_INVALID）
- 并发 >3 排队（status=pending，由 alerting_worker 队列执行）；60s 超时转异步 + 站内通知
- 强制水印"由 AgendaScope 观澜生成" + 数据口径声明；版权合规 L1：仅标题 + ≤150 字摘录，不含正文全文
- 数据全部来自 topics/agenda_snapshots/articles 真实查询；文件保留 7 天

请求契约（双兼容，前端 reportExports.ts 与外部对接约定）：
  {template|report_type, format: pdf|docx, scope|params: {topic_id?, countries?, from, to}, locale?}
  time_range 亦可单独传 {from, to}（覆盖 scope 内 from/to）。
"""
from __future__ import annotations

import threading
import time
import uuid
from datetime import UTC, date, datetime, timedelta
from pathlib import Path
from typing import Any

from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.errors import CODE_NOT_FOUND, CODE_PARAM_INVALID, BizError
from app.core.logging import get_logger
from app.models.alert import Alert, AlertRule
from app.models.article import Article
from app.models.report import ReportExport
from app.models.topic import AgendaSnapshot, Topic, TopicArticle

logger = get_logger("services.report")

MAX_WINDOW_DAYS = 90
MAX_CONCURRENT_EXPORTS = 3
SYNC_TIMEOUT_SECONDS = 60
EXPORT_TTL_DAYS = 7
EXCERPT_MAX = 150
_TOP_ARTICLES = 20
_TOP_ROWS = 30

DEFAULT_EXPORT_DIR = "data/report_exports"
WATERMARK = "由 AgendaScope 观澜生成"
DATA_DISCLAIMER = (
    "数据口径声明：本报告数据来自 AgendaScope 观澜平台监测范围内的公开媒体报道，"
    "显著性/情感指标为平台算法测算结果，统计关联不等于因果关系；"
    "报道内容版权归原媒体所有，本报告仅含标题与不超过 150 字摘录，不含正文全文。"
)

TEMPLATES = ("topic_deep", "compare_brief", "periodic_weekly")
FORMATS = ("pdf", "docx")

_TEMPLATE_NAMES = {
    "topic_deep": "议题深度报告",
    "compare_brief": "跨国对比简报",
    "periodic_weekly": "周期监测周报",
}


# ---------------------------------------------------------------------------
# 请求契约归一化与校验
# ---------------------------------------------------------------------------


def _parse_date(value: Any, field: str) -> date:
    if isinstance(value, date) and not isinstance(value, datetime):
        return value
    if isinstance(value, datetime):
        return value.date()
    try:
        return date.fromisoformat(str(value)[:10])
    except ValueError:
        raise BizError(CODE_PARAM_INVALID, f"{field} 日期格式非法（应 YYYY-MM-DD）: {value}") from None


def normalize_export_payload(payload: dict[str, Any]) -> dict[str, Any]:
    """归一化 {template, format, scope{topic_id,countries,from,to}, locale} 并校验。

    兼容字段：report_type→template，params→scope，time_range({from,to}|{start,end})→scope.from/to。
    """
    template = payload.get("template") or payload.get("report_type")
    if template not in TEMPLATES:
        raise BizError(CODE_PARAM_INVALID, f"template 仅支持 {list(TEMPLATES)}")

    fmt = payload.get("format")
    if fmt not in FORMATS:
        raise BizError(CODE_PARAM_INVALID, f"format 仅支持 {list(FORMATS)}")

    scope = dict(payload.get("scope") or payload.get("params") or {})
    time_range = payload.get("time_range") or {}
    if time_range:
        scope["from"] = time_range.get("from") or time_range.get("start") or scope.get("from")
        scope["to"] = time_range.get("to") or time_range.get("end") or scope.get("to")

    date_from = _parse_date(scope.get("from"), "from")
    date_to = _parse_date(scope.get("to"), "to")
    if date_to < date_from:
        raise BizError(CODE_PARAM_INVALID, "时间窗终点早于起点")
    if (date_to - date_from).days > MAX_WINDOW_DAYS:
        raise BizError(
            CODE_PARAM_INVALID,
            f"时间窗超过 {MAX_WINDOW_DAYS} 天上限（{(date_to - date_from).days} 天）",
        )

    countries = [str(c).upper() for c in (scope.get("countries") or [])]
    topic_id = scope.get("topic_id")
    if template == "topic_deep" and not topic_id:
        raise BizError(CODE_PARAM_INVALID, "topic_deep 模板需提供 scope.topic_id")
    if template == "compare_brief" and not (2 <= len(countries) <= 4):
        raise BizError(CODE_PARAM_INVALID, "compare_brief 模板需提供 2-4 个国家")

    return {
        "template": template,
        "format": fmt,
        "scope": {
            "topic_id": str(topic_id) if topic_id else None,
            "countries": countries,
            "from": date_from.isoformat(),
            "to": date_to.isoformat(),
        },
        "locale": payload.get("locale") or "zh-CN",
    }


def scope_summary(template: str, scope: dict, topic_name: str | None = None) -> str:
    parts = []
    if template == "topic_deep":
        parts.append(topic_name or str(scope.get("topic_id") or ""))
    elif scope.get("countries"):
        parts.append(",".join(scope["countries"]))
    else:
        parts.append("全球")
    parts.append(f"{scope.get('from','')}~{scope.get('to','')}")
    return " / ".join(p for p in parts if p)


# ---------------------------------------------------------------------------
# 数据构建（真实查询；输出统一 report dict 供 PDF/DOCX 渲染）
# ---------------------------------------------------------------------------


def _window(scope: dict) -> tuple[datetime, datetime]:
    start = datetime.fromisoformat(scope["from"]).replace(tzinfo=UTC)
    end = (datetime.fromisoformat(scope["to"]) + timedelta(days=1)).replace(tzinfo=UTC)
    return start, end


def _base_report(template: str, scope: dict, title: str, summary: str) -> dict[str, Any]:
    return {
        "title": title,
        "template": template,
        "template_name": _TEMPLATE_NAMES[template],
        "generated_at": datetime.now(UTC).isoformat(timespec="seconds"),
        "watermark": WATERMARK,
        "disclaimer": DATA_DISCLAIMER,
        "scope_summary": summary,
        "sections": [],
    }


def _article_excerpt_rows(db: Session, topic_id: uuid.UUID, start: datetime, end: datetime) -> list[list[str]]:
    """代表报道：标题 + ≤150 字摘录 + 原文链接（版权合规 L1，不出正文全文）。"""
    stmt = (
        select(Article)
        .join(TopicArticle, TopicArticle.article_id == Article.id)
        .where(TopicArticle.topic_id == topic_id, Article.published_at >= start, Article.published_at <= end)
        .order_by(Article.published_at.desc())
        .limit(_TOP_ARTICLES)
    )
    rows = []
    for a in db.scalars(stmt).all():
        excerpt = (a.content or "")[:EXCERPT_MAX].strip()
        rows.append([
            a.published_at.strftime("%Y-%m-%d") if a.published_at else "",
            a.title or "",
            excerpt,
            a.url or "",
        ])
    return rows


def build_topic_deep(db: Session, scope: dict) -> dict[str, Any]:
    """议题深度报告：概览 + 分国显著性/报道量 + 情感 + 代表报道摘录。"""
    topic = db.get(Topic, uuid.UUID(scope["topic_id"]))
    if topic is None:
        raise BizError(CODE_NOT_FOUND, f"议题不存在: {scope['topic_id']}")
    start, end = _window(scope)

    report = _base_report(
        "topic_deep", scope,
        f"议题深度报告：{topic.name_zh or topic.name}",
        scope_summary("topic_deep", scope, topic.name_zh or topic.name),
    )
    report["sections"].append({
        "heading": "一、议题概览",
        "paragraphs": [
            f"议题名称：{topic.name_zh or topic.name}（{topic.name_auto}）",
            f"议题分类：{topic.topic_category or '未分类'}",
            f"生命周期：{topic.lifecycle_state} / 置信度：{topic.confidence}",
            f"关键词：{'、'.join(topic.keywords or []) or '无'}",
            f"摘要：{topic.summary_zh or '暂无摘要'}",
            f"统计窗口：{scope['from']} ~ {scope['to']}",
        ],
        "table": None,
    })

    snap_stmt = (
        select(AgendaSnapshot)
        .where(
            AgendaSnapshot.topic_id == topic.id,
            AgendaSnapshot.window_end >= start, AgendaSnapshot.window_end <= end,
        )
        .order_by(AgendaSnapshot.country_code, AgendaSnapshot.window_end)
    )
    snaps = list(db.scalars(snap_stmt).all())
    by_country: dict[str, dict] = {}
    for s in snaps:
        agg = by_country.setdefault(s.country_code, {"articles": 0, "best_rank": None, "neg": [], "pos": []})
        agg["articles"] += int(s.article_count)
        if agg["best_rank"] is None or int(s.salience_rank) < agg["best_rank"]:
            agg["best_rank"] = int(s.salience_rank)
        if s.sentiment_neg is not None:
            agg["neg"].append(float(s.sentiment_neg))
        if s.sentiment_pos is not None:
            agg["pos"].append(float(s.sentiment_pos))
    rows = [
        [cc, str(a["articles"]), str(a["best_rank"] or "-"),
         f"{sum(a['neg'])/len(a['neg']):.2f}" if a["neg"] else "-",
         f"{sum(a['pos'])/len(a['pos']):.2f}" if a["pos"] else "-"]
        for cc, a in sorted(by_country.items(), key=lambda kv: -kv[1]["articles"])[:_TOP_ROWS]
    ]
    report["sections"].append({
        "heading": "二、分国报道量与显著性",
        "paragraphs": [] if rows else ["窗口内该议题无快照数据。"],
        "table": {"headers": ["国家", "报道量", "最佳显著性排名", "负面情感占比(均值)", "正面情感占比(均值)"], "rows": rows} if rows else None,
    })

    article_rows = _article_excerpt_rows(db, topic.id, start, end)
    report["sections"].append({
        "heading": "三、代表报道（标题与摘录）",
        "paragraphs": [] if article_rows else ["窗口内无关联报道。"],
        "table": {"headers": ["日期", "标题", "摘录(≤150字)", "原文链接"], "rows": article_rows} if article_rows else None,
    })
    return report


def build_compare_brief(db: Session, scope: dict) -> dict[str, Any]:
    """跨国对比简报：2-4 国 Top 议题 + 情感对比。"""
    start, end = _window(scope)
    countries = scope["countries"]
    report = _base_report(
        "compare_brief", scope,
        f"跨国对比简报：{' vs '.join(countries)}",
        scope_summary("compare_brief", scope),
    )
    contrast_rows: list[list[str]] = []
    for cc in countries:
        stmt = (
            select(AgendaSnapshot, Topic)
            .join(Topic, Topic.id == AgendaSnapshot.topic_id)
            .where(
                AgendaSnapshot.country_code == cc,
                AgendaSnapshot.window_end >= start, AgendaSnapshot.window_end <= end,
            )
            .order_by(AgendaSnapshot.salience_rank.asc())
            .limit(10)
        )
        rows = db.execute(stmt).all()
        topic_names: dict[str, str] = {}
        negs: list[float] = []
        total_articles = 0
        table_rows: list[list[str]] = []
        seen: set[str] = set()
        for snap, topic in rows:
            key = str(topic.id)
            if key in seen:
                continue
            seen.add(key)
            topic_names[key] = topic.name_zh or topic.name
            total_articles += int(snap.article_count)
            if snap.sentiment_neg is not None:
                negs.append(float(snap.sentiment_neg))
            table_rows.append([
                str(snap.salience_rank), topic.name_zh or topic.name,
                str(int(snap.article_count)),
                f"{float(snap.sentiment_neg):.2f}" if snap.sentiment_neg is not None else "-",
            ])
        contrast_rows.append([
            cc, str(total_articles), str(len(table_rows)),
            f"{sum(negs)/len(negs):.2f}" if negs else "-",
        ])
        report["sections"].append({
            "heading": f"{'一二三四'[countries.index(cc)]}、{cc} Top 议题",
            "paragraphs": [] if table_rows else [f"窗口内 {cc} 无快照数据（该国本期数据缺失，不代表无报道）。"],
            "table": {"headers": ["排名", "议题", "报道量", "负面情感占比"], "rows": table_rows} if table_rows else None,
        })
    report["sections"].append({
        "heading": "对比小结",
        "paragraphs": ["统计关联不等于因果，跨国报道量差异受媒体覆盖与语言因素影响。"],
        "table": {"headers": ["国家", "Top议题报道总量", "上榜议题数", "负面情感占比(均值)"], "rows": contrast_rows},
    })
    return report


def build_periodic_weekly(db: Session, scope: dict) -> dict[str, Any]:
    """周期监测周报：窗口内分国 Top 议题总览。"""
    start, end = _window(scope)
    report = _base_report(
        "periodic_weekly", scope,
        f"周期监测周报：{scope['from']} ~ {scope['to']}",
        scope_summary("periodic_weekly", scope),
    )
    stmt = (
        select(AgendaSnapshot, Topic)
        .join(Topic, Topic.id == AgendaSnapshot.topic_id)
        .where(AgendaSnapshot.window_end >= start, AgendaSnapshot.window_end <= end)
        .order_by(AgendaSnapshot.country_code, AgendaSnapshot.salience_rank.asc())
    )
    if scope.get("countries"):
        stmt = stmt.where(AgendaSnapshot.country_code.in_(scope["countries"]))
    by_country: dict[str, list[list[str]]] = {}
    seen: set[tuple[str, str]] = set()
    for snap, topic in db.execute(stmt).all():
        key = (snap.country_code, str(topic.id))
        if key in seen:
            continue
        seen.add(key)
        lst = by_country.setdefault(snap.country_code, [])
        if len(lst) >= 5:
            continue
        lst.append([
            str(snap.salience_rank), topic.name_zh or topic.name,
            str(int(snap.article_count)),
            f"{float(snap.sentiment_neg):.2f}" if snap.sentiment_neg is not None else "-",
        ])
    if not by_country:
        report["sections"].append({
            "heading": "监测总览",
            "paragraphs": ["窗口内监测范围无快照数据（数据缺失，不代表无报道）。"],
            "table": None,
        })
        return report
    for cc, rows in sorted(by_country.items()):
        report["sections"].append({
            "heading": f"{cc} Top 议题（{len(rows)}）",
            "paragraphs": [],
            "table": {"headers": ["排名", "议题", "报道量", "负面情感占比"], "rows": rows},
        })
    return report


_BUILDERS = {
    "topic_deep": build_topic_deep,
    "compare_brief": build_compare_brief,
    "periodic_weekly": build_periodic_weekly,
}


# ---------------------------------------------------------------------------
# 渲染（PDF / DOCX；每页/页脚强制水印 + 数据口径声明）
# ---------------------------------------------------------------------------


def render_pdf(report: dict[str, Any], dest_path: str | Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    font = "STSong-Light"
    body = ParagraphStyle("body", fontName=font, fontSize=10, leading=15)
    h1 = ParagraphStyle("h1", fontName=font, fontSize=18, leading=24, spaceAfter=8)
    h2 = ParagraphStyle("h2", fontName=font, fontSize=13, leading=18, spaceBefore=10, spaceAfter=4)
    small = ParagraphStyle("small", fontName=font, fontSize=8, leading=11, textColor="#666666")

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFont(font, 8)
        canvas.setFillGray(0.5)
        canvas.drawString(15 * mm, 10 * mm, f"{report['watermark']} ｜ 生成时间 {report['generated_at']}")
        canvas.drawCentredString(A4[0] / 2, A4[1] - 12 * mm, report["watermark"])
        canvas.restoreState()

    doc = SimpleDocTemplate(str(dest_path), pagesize=A4, title=report["title"])
    story: list = [
        Paragraph(report["title"], h1),
        Paragraph(f"模板：{report['template_name']} ｜ 范围：{report['scope_summary']}", small),
        Paragraph(report["disclaimer"], small),
        Spacer(1, 6),
    ]
    for section in report["sections"]:
        story.append(Paragraph(section["heading"], h2))
        for para in section.get("paragraphs") or []:
            story.append(Paragraph(str(para).replace("\n", "<br/>"), body))
        table = section.get("table")
        if table and table.get("rows"):
            data = [[Paragraph(str(h), body) for h in table["headers"]]]
            for row in table["rows"]:
                data.append([Paragraph(str(c), small) for c in row])
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.25, "#999999"),
                ("BACKGROUND", (0, 0), (-1, 0), "#EEEEEE"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(t)
        story.append(Spacer(1, 4))
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


def render_docx(report: dict[str, Any], dest_path: str | Path) -> None:
    import docx

    document = docx.Document()
    # 页脚强制水印 + 口径声明
    footer = document.sections[0].footer.paragraphs[0]
    footer.text = f"{report['watermark']} ｜ {report['disclaimer']}"

    document.add_heading(report["title"], level=0)
    document.add_paragraph(f"模板：{report['template_name']} ｜ 范围：{report['scope_summary']}")
    document.add_paragraph(report["disclaimer"])
    for section in report["sections"]:
        document.add_heading(section["heading"], level=1)
        for para in section.get("paragraphs") or []:
            document.add_paragraph(str(para))
        table = section.get("table")
        if table and table.get("rows"):
            t = document.add_table(rows=1, cols=len(table["headers"]))
            t.style = "Table Grid"
            for i, h in enumerate(table["headers"]):
                t.rows[0].cells[i].text = str(h)
            for row in table["rows"]:
                cells = t.add_row().cells
                for i, c in enumerate(row):
                    cells[i].text = str(c)
    document.save(str(dest_path))


_RENDERERS = {"pdf": render_pdf, "docx": render_docx}


# ---------------------------------------------------------------------------
# 生成执行（同步核心，供内联 60s 尝试与 worker 队列复用）
# ---------------------------------------------------------------------------


def generate_export(
    db: Session,
    export: ReportExport,
    export_dir: str | Path = DEFAULT_EXPORT_DIR,
) -> ReportExport:
    """执行生成：构建数据 → 渲染文件 → 更新行状态（done/failed）。"""
    started = time.monotonic()
    export.status = "processing"
    db.flush()
    try:
        report = _BUILDERS[export.template](db, dict(export.scope or {}))
        out_dir = Path(export_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        dest = out_dir / f"{export.id}.{export.format}"
        _RENDERERS[export.format](report, dest)
        export.file_path = str(dest)
        export.file_size = dest.stat().st_size
        export.status = "done"
        export.error = None
        export.expires_at = datetime.now(UTC) + timedelta(days=EXPORT_TTL_DAYS)
    except BizError:
        raise
    except Exception as exc:  # noqa: BLE001 失败真实落库，不伪造成功
        export.status = "failed"
        export.error = str(exc)[:500]
        logger.error("report_export_fail", export_id=str(export.id), error=str(exc)[:300])
    export.duration_ms = int((time.monotonic() - started) * 1000)
    db.flush()
    return export


def _ensure_notify_rule(db: Session) -> AlertRule:
    from app.services.seed_service import ensure_admin

    admin = ensure_admin(db)
    rule = db.scalar(select(AlertRule).where(AlertRule.name == "系统-报告导出通知"))
    if rule is None:
        rule = AlertRule(
            user_id=admin.id,
            name="系统-报告导出通知",
            country_codes=[],
            keywords=["__report_export__"],
            condition_type="growth_rate",
            condition_value=0,
            notify_channels=["inapp"],
        )
        db.add(rule)
        db.flush()
    return rule


def notify_export_finished(db: Session, export: ReportExport) -> None:
    """异步完成后的站内通知（60s 超时转异步 / 队列完成的任务）。"""
    rule = _ensure_notify_rule(db)
    done = export.status == "done"
    alert = Alert(
        rule_id=rule.id,
        user_id=export.user_id,
        payload={
            "kind": "report_export_done" if done else "report_export_failed",
            "message": (
                f"您的{_TEMPLATE_NAMES.get(export.template, export.template)}已生成，可前往报告中心下载。"
                if done else
                f"您的{_TEMPLATE_NAMES.get(export.template, export.template)}生成失败：{export.error or '未知错误'}"
            ),
            "export_id": str(export.id),
            "download_url": f"/api/v1/report-exports/{export.id}/download" if done else None,
        },
        status="unread",
    )
    db.add(alert)
    db.flush()


def count_active_exports(db: Session) -> int:
    return int(db.scalar(
        select(func.count()).select_from(ReportExport).where(
            ReportExport.status.in_(["pending", "processing"]),
        )
    ) or 0)


def create_export(db: Session, user_id: uuid.UUID, payload: dict[str, Any]) -> ReportExport:
    """创建导出任务（校验 + 落库 pending；不执行生成）。"""
    normalized = normalize_export_payload(payload)
    export = ReportExport(
        user_id=user_id,
        template=normalized["template"],
        format=normalized["format"],
        scope=normalized["scope"],
        locale=normalized["locale"],
        status="pending",
        watermark=f"{WATERMARK} + 数据口径声明",
    )
    db.add(export)
    db.flush()
    return export


def process_pending_exports(
    db: Session,
    export_dir: str | Path = DEFAULT_EXPORT_DIR,
) -> int:
    """worker 队列执行：pending 任务按创建序生成（单轮至多并发上限个）。返回处理数。"""
    waiting = list(db.scalars(
        select(ReportExport)
        .where(ReportExport.status == "pending")
        .order_by(ReportExport.created_at.asc())
        .limit(max(MAX_CONCURRENT_EXPORTS, 1))
    ).all())
    processed = 0
    for export in waiting[: max(MAX_CONCURRENT_EXPORTS, 1)]:
        try:
            generate_export(db, export, export_dir)
            notify_export_finished(db, export)
            db.flush()
            processed += 1
        except BizError as exc:
            export.status = "failed"
            export.error = exc.message[:500]
            db.flush()
            processed += 1
        except Exception as exc:  # noqa: BLE001 单个任务失败不阻塞队列
            db.rollback()
            logger.error("report_export_queue_item_fail", export_id=str(export.id), error=str(exc)[:300])
    return processed


class InlineRunResult:
    """内联执行协调：路由 join 超时后将任务标记为异步，完成时站内通知。"""

    def __init__(self):
        self.async_mode = False
        self.export: ReportExport | None = None
        self.error: Exception | None = None


def run_inline_with_timeout(
    export_id: uuid.UUID,
    session_factory,
    export_dir: str | Path = DEFAULT_EXPORT_DIR,
    timeout_seconds: int = SYNC_TIMEOUT_SECONDS,
) -> InlineRunResult:
    """60s 预算内联生成：超时转异步（线程继续，完成后站内通知）。"""
    result = InlineRunResult()

    def target() -> None:
        db = session_factory()
        try:
            export = db.get(ReportExport, export_id)
            if export is None or export.status == "done":
                result.export = export
                return
            generate_export(db, export, export_dir)
            if result.async_mode:
                notify_export_finished(db, export)
            db.commit()
            result.export = export
        except BizError as exc:
            db.rollback()
            db2 = session_factory()
            try:
                export = db2.get(ReportExport, export_id)
                if export is not None:
                    export.status = "failed"
                    export.error = exc.message[:500]
                    db2.commit()
                    result.export = export
            finally:
                db2.close()
        except Exception as exc:  # noqa: BLE001
            db.rollback()
            result.error = exc
            logger.error("report_export_inline_fail", export_id=str(export_id), error=str(exc)[:300])
        finally:
            db.close()

    thread = threading.Thread(target=target, daemon=True, name=f"report-export-{export_id}")
    thread.start()
    thread.join(timeout_seconds)
    if thread.is_alive():
        result.async_mode = True  # 超时转异步：线程完成后由 notify_export_finished 通知
    return result


__all__ = [
    "DATA_DISCLAIMER",
    "DEFAULT_EXPORT_DIR",
    "EXPORT_TTL_DAYS",
    "FORMATS",
    "MAX_CONCURRENT_EXPORTS",
    "MAX_WINDOW_DAYS",
    "SYNC_TIMEOUT_SECONDS",
    "TEMPLATES",
    "WATERMARK",
    "build_compare_brief",
    "build_periodic_weekly",
    "build_topic_deep",
    "count_active_exports",
    "create_export",
    "generate_export",
    "normalize_export_payload",
    "notify_export_finished",
    "process_pending_exports",
    "render_docx",
    "render_pdf",
    "run_inline_with_timeout",
    "scope_summary",
]
