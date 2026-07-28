"""report_service 单元测试（T4.17）：契约归一化 / 90 天预检 / PDF/DOCX 渲染（真实生成文件）。"""
from datetime import date, timedelta

import pytest

from app.core.errors import BizError
from app.services import report_service
from app.services.report_service import (
    DATA_DISCLAIMER,
    MAX_WINDOW_DAYS,
    WATERMARK,
    normalize_export_payload,
    render_docx,
    render_pdf,
    scope_summary,
)

FROM = date(2026, 7, 1).isoformat()
TO = date(2026, 7, 24).isoformat()


class TestNormalizePayload:
    def test_template_contract(self):
        out = normalize_export_payload({
            "template": "topic_deep", "format": "pdf",
            "scope": {"topic_id": "3f2a0000-0000-0000-0000-000000000001", "from": FROM, "to": TO},
        })
        assert out["template"] == "topic_deep" and out["format"] == "pdf"
        assert out["scope"]["from"] == FROM and out["scope"]["to"] == TO

    def test_report_type_params_time_range_contract(self):
        """任务书契约：report_type + params + time_range 等价归一。"""
        out = normalize_export_payload({
            "report_type": "compare_brief", "format": "docx",
            "params": {"countries": ["us", "gb"]},
            "time_range": {"from": FROM, "to": TO},
        })
        assert out["template"] == "compare_brief"
        assert out["scope"]["countries"] == ["US", "GB"]
        assert out["scope"]["from"] == FROM and out["scope"]["to"] == TO

    def test_window_over_90_days_rejected(self):
        far_to = (date(2026, 7, 1) + timedelta(days=MAX_WINDOW_DAYS + 1)).isoformat()
        with pytest.raises(BizError) as exc:
            normalize_export_payload({
                "template": "periodic_weekly", "format": "pdf",
                "scope": {"from": FROM, "to": far_to},
            })
        assert exc.value.code == 1001

    def test_window_exactly_90_days_ok(self):
        to90 = (date(2026, 7, 1) + timedelta(days=MAX_WINDOW_DAYS)).isoformat()
        out = normalize_export_payload({
            "template": "periodic_weekly", "format": "pdf",
            "scope": {"from": FROM, "to": to90},
        })
        assert out["scope"]["to"] == to90

    def test_bad_template_rejected(self):
        with pytest.raises(BizError):
            normalize_export_payload({"template": "nope", "format": "pdf", "scope": {"from": FROM, "to": TO}})

    def test_bad_format_rejected(self):
        with pytest.raises(BizError):
            normalize_export_payload({"template": "periodic_weekly", "format": "xlsx", "scope": {"from": FROM, "to": TO}})

    def test_topic_deep_requires_topic_id(self):
        with pytest.raises(BizError):
            normalize_export_payload({"template": "topic_deep", "format": "pdf", "scope": {"from": FROM, "to": TO}})

    def test_compare_brief_country_count(self):
        with pytest.raises(BizError):
            normalize_export_payload({
                "template": "compare_brief", "format": "pdf",
                "scope": {"countries": ["US"], "from": FROM, "to": TO},
            })

    def test_reversed_window_rejected(self):
        with pytest.raises(BizError):
            normalize_export_payload({
                "template": "periodic_weekly", "format": "pdf",
                "scope": {"from": TO, "to": FROM},
            })


class TestScopeSummary:
    def test_with_countries(self):
        s = scope_summary("compare_brief", {"countries": ["US", "GB"], "from": FROM, "to": TO})
        assert "US,GB" in s and FROM in s

    def test_topic_deep_uses_name(self):
        s = scope_summary("topic_deep", {"from": FROM, "to": TO}, "新疆棉争议")
        assert s.startswith("新疆棉争议")


def _sample_report() -> dict:
    return {
        "title": "议题深度报告：测试议题",
        "template": "topic_deep",
        "template_name": "议题深度报告",
        "generated_at": "2026-07-28T12:00:00+00:00",
        "watermark": WATERMARK,
        "disclaimer": DATA_DISCLAIMER,
        "scope_summary": "测试议题 / 2026-07-01~2026-07-24",
        "sections": [
            {"heading": "一、议题概览", "paragraphs": ["议题名称：测试议题", "摘要：中文内容渲染验证"], "table": None},
            {
                "heading": "二、分国报道量",
                "paragraphs": [],
                "table": {
                    "headers": ["国家", "报道量", "排名"],
                    "rows": [["US", "42", "1"], ["JP", "17", "2"]],
                },
            },
        ],
    }


class TestRenderers:
    def test_render_pdf_real_file(self, tmp_path):
        dest = tmp_path / "r.pdf"
        render_pdf(_sample_report(), dest)
        data = dest.read_bytes()
        assert data.startswith(b"%PDF") and len(data) > 500

    def test_render_docx_real_file(self, tmp_path):
        dest = tmp_path / "r.docx"
        render_docx(_sample_report(), dest)
        import docx

        document = docx.Document(str(dest))
        texts = [p.text for p in document.paragraphs]
        assert any("测试议题" in t for t in texts)
        footer = document.sections[0].footer.paragraphs[0].text
        assert WATERMARK in footer and "数据口径声明" in footer
